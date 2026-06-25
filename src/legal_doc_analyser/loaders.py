"""
Load document text from a file (.txt, .md, .pdf, .docx) or standard input.

Real legal documents arrive as PDFs and Word files, not plain text. These
loaders turn those formats into the raw text the analyser works on. Image-only
("scanned") PDFs are detected and reported honestly: OCR is out of scope for this
prototype, so the tool says so rather than returning silent gibberish.
"""

from __future__ import annotations

import sys
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class DocumentLoadError(Exception):
    """Raised when a document cannot be read or contains no extractable text."""


def _load_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages).strip()
    if not text:
        raise DocumentLoadError(
            f"No extractable text in '{path.name}'. It may be a scanned image; "
            "this prototype does not perform OCR."
        )
    return text


def _load_docx(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts: list[str] = [para.text for para in document.paragraphs]
    # Word tables hold a lot of contractual content; include their cells too.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(p for p in parts if p and p.strip()).strip()
    if not text:
        raise DocumentLoadError(f"No extractable text in '{path.name}'.")
    return text


def load_from_path(path: Path) -> str:
    """Load document text from a file, dispatching on its extension."""
    if not path.exists():
        raise DocumentLoadError(f"File not found: {path}")
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        return _load_plain(path)
    if ext == ".pdf":
        return _load_pdf(path)
    if ext == ".docx":
        return _load_docx(path)
    raise DocumentLoadError(
        f"Unsupported file type '{ext}'. Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
    )


def load_from_stdin() -> str:
    """Read document text piped via standard input."""
    return sys.stdin.read()


def find_documents(directory: Path) -> list[Path]:
    """Return supported document files in a directory, sorted by name."""
    if not directory.is_dir():
        raise DocumentLoadError(f"Not a directory: {directory}")
    return sorted(
        p for p in directory.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )
